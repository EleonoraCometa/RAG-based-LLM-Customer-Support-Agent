"""
AI Agent
Architettura: LangChain Agent + ChromaDB + Tool calling
LLM: Google Gemini 2.5 Flash 
"""

import os
import pandas as pd
from dotenv import load_dotenv
from docx import Document

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings, HarmBlockThreshold, HarmCategory
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
from langchain.memory import ConversationBufferWindowMemory
from langchain.agents import AgentExecutor, create_tool_calling_agent
from langchain.tools import Tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder

load_dotenv() 

# ─────────────────────────────────────────────
# 1. CARICAMENTO DATI
# ─────────────────────────────────────────────

DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data") 


def load_tone_of_voice() -> str:
    path = os.path.join(DATA_DIR, "tone_of_voice.txt")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_tax_knowledge() -> list[LCDocument]:
    """Carica i due file .docx della knowledge base fiscale."""
    docs = []
    for fname in ["tax_knowledge_1.docx", "tax_knowledge_2.docx"]:
        path = os.path.join(DATA_DIR, fname)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        docs.append(LCDocument(page_content=text, metadata={"source": fname}))
    return docs


def load_customer_data() -> pd.DataFrame:
    path = os.path.join(DATA_DIR, "customer_data.xlsx")
    df = pd.read_excel(path)
    
    for col in ["apertura_piva"]: 
        if col in df.columns:
            try:
                df[col] = pd.to_datetime(df[col], errors="coerce").dt.strftime("%Y-%m-%d")
            except Exception:
                df[col] = df[col].astype(str)
    return df


# ─────────────────────────────────────────────
# 2. VECTOR STORE (RAG)
# ─────────────────────────────────────────────

def build_vector_store(docs: list[LCDocument]) -> Chroma:
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=80)
    chunks = splitter.split_documents(docs)
    embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
    vs = Chroma.from_documents(chunks, embeddings, collection_name="tax_knowledge")
    return vs


# ─────────────────────────────────────────────
# 3. TOOLS
# ─────────────────────────────────────────────

_last_tools_used = []


def reset_tool_tracking():
    global _last_tools_used
    _last_tools_used = []


def get_tools_used():
    return list(_last_tools_used)


def make_rag_tool(vector_store: Chroma):
    retriever = vector_store.as_retriever(search_kwargs={"k": 4})

    def rag_query(query: str) -> str:
        _last_tools_used.append("RAG")
        results = retriever.invoke(query)
        if not results:
            return "Nessuna informazione trovata nella knowledge base fiscale."
        context = "\n\n---\n\n".join(r.page_content for r in results)
        return f"Informazioni trovate nella knowledge base:\n\n{context}"

    return Tool(
        name="consulta_normativa_fiscale",
        func=rag_query,
        description=(
            "Usa questo tool per rispondere a domande su tasse, regimi fiscali (forfettario, semplificato), "
            "detrazioni, fatturazione elettronica, note di credito e scadenze fiscali. "
            "Input: la domanda del cliente in italiano."
        ),
    )


def make_customer_tool(df: pd.DataFrame, current_client_name: dict):
    """current_client_name è un dict mutabile per passare il cliente attivo."""

    def get_customer_info(query: str = "") -> str:
        _last_tools_used.append("CUSTOMER")
        nome = current_client_name.get("name")
        if not nome:
            return "Nessun cliente autenticato. Chiedi al cliente di accedere alla piattaforma."

        nome_lower = nome.strip().lower()
        parts = nome_lower.split()
        mask = pd.Series([False] * len(df))
        for part in parts:
            mask |= df["nome"].str.lower().str.contains(part, na=False)
            mask |= df["cognome"].str.lower().str.contains(part, na=False)
        results = df[mask]
        if results.empty:
            return f"Cliente '{nome}' non trovato nel database."
        row = results.iloc[0]
        return (
            f"Dati cliente attualmente autenticato:\n"
            f"- Nome: {row['nome']} {row['cognome']}\n"
            f"- Regime fiscale: {row['regime']}\n"
            f"- Cassa previdenziale: {row['cassa']}\n"
            f"- Commercialista assegnato: {row['commercialista']}\n"
            f"- Apertura P.IVA: {row['apertura_piva']}\n"
            f"- Fatturato 2025: €{row['fatturato_2025']}k\n"
            f"- Fatturato 2026 (YTD): €{row['fatturato_2026']}k"
        )

    return Tool(
        name="consulta_dati_cliente",
        func=get_customer_info,
        description=(
            "Usa questo tool quando il cliente fa domande sui propri dati personali o fiscali: "
            "regime fiscale, fatturato, cassa previdenziale, commercialista assegnato, data apertura P.IVA. "
            "Non richiede input: il cliente è già autenticato e i dati vengono recuperati automaticamente."
        ),
    )


def make_escalation_tool():
    def escalate(motivo: str) -> str:
        _last_tools_used.append("ESCALATION")
        return (
            f"ESCALATION_RICHIESTA: {motivo}\n\n"
            "Comunica al cliente che lo metterai in contatto con un Customer Success Consultant umano."
        )

    return Tool(
        name="escalation_customer_success",
        func=escalate,
        description=(
            "Usa questo tool quando: (1) la domanda è troppo specifica o complessa per una risposta automatica, "
            "(2) il cliente esprime insoddisfazione, (3) la situazione richiede un'analisi personalizzata approfondita, "
            "(4) non hai abbastanza informazioni per rispondere in modo affidabile, "
            "(5) il cliente chiede esplicitamente di parlare con un consulente. "
            "Input: breve descrizione del motivo dell'escalation."
        ),
    )


# ─────────────────────────────────────────────
# 4. PROMPT DI SISTEMA
# ─────────────────────────────────────────────

SYSTEM_PROMPT = """Sei l'AI Customer Support Agent. Rispondi ai clienti con partita IVA.

TONO DI VOCE (rispettalo sempre):
{tone_of_voice}

REGOLE DI COMPORTAMENTO:
- Usa sempre "noi" per riferirti a noi, mai "io".
- Frasi brevi, massimo 20 parole per frase.
- Spiega i termini tecnici la prima volta che li usi.
- Sii diretto: usa l'imperativo per i passaggi operativi.
- Se non sei sicuro, usa il tool di escalation invece di inventare.

QUANDO USARE I TOOL:
- Se il cliente fa una domanda sui suoi dati personali (regime, fatturato, commercialista) → usa consulta_dati_cliente
- Se il cliente fa una domanda fiscale o sulla piattaforma → usa consulta_normativa_fiscale
- Se la domanda è fuori scope, complessa, o il cliente chiede aiuto umano → usa escalation_customer_success

Dopo aver usato i tool, formula la risposta finale al cliente con questo tono.
"""


# ─────────────────────────────────────────────
# 5. AGENTE PRINCIPALE
# ─────────────────────────────────────────────

class CustomerSupportAgent:
    def __init__(self):
        # Safety settings: disabilitiamo i filtri Gemini per evitare blocchi falsi positivi
        # su contenuti fiscali innocui (es. "fattura", "tasse" possono essere flaggati per errore)
        safety_settings = {
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
        }
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            temperature=0.2,
            safety_settings=safety_settings,
        )
        self.tone = load_tone_of_voice()
        self.df = load_customer_data()

        tax_docs = load_tax_knowledge()
        self.vector_store = build_vector_store(tax_docs)

        self.active_client = {"name": None}

        tools = [
            make_rag_tool(self.vector_store),
            make_customer_tool(self.df, self.active_client),
            make_escalation_tool(),
        ]

        # Memory windowed: tiene solo gli ultimi 4 scambi.
        self.memory = ConversationBufferWindowMemory(
            k=4,
            memory_key="chat_history",
            return_messages=True,
        )

        prompt = ChatPromptTemplate.from_messages([
            ("system", SYSTEM_PROMPT.format(tone_of_voice=self.tone)),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}"),
            MessagesPlaceholder(variable_name="agent_scratchpad"),
        ])

        agent = create_tool_calling_agent(self.llm, tools, prompt)
        self.executor = AgentExecutor(
            agent=agent,
            tools=tools,
            memory=self.memory,
            verbose=True,
            max_iterations=8,
            handle_parsing_errors=True,
            return_intermediate_steps=False,
        )

    def set_active_client(self, nome: str | None):
        """Imposta il cliente autenticato (simula login)."""
        self.active_client["name"] = nome

    def reset_memory(self):
        self.memory.clear()

    def ask(self, domanda: str, nome_cliente: str | None = None) -> dict:
        """
        Processa una domanda del cliente.
        Restituisce: risposta, tool usato, escalation flag, rate_limited flag.
        """
        import time
        import re

        if nome_cliente is not None:
            self.set_active_client(nome_cliente)

        reset_tool_tracking()

        if self.active_client["name"]:
            input_text = f"[Cliente autenticato: {self.active_client['name']}] {domanda}"
        else:
            input_text = domanda

        # Tre stati possibili:
        # 1. Successo → output con testo
        # 2. Rate limit (429) → messaggio dedicato per l'utente
        # 3. Risposta vuota → escalation
        output = ""
        rate_limited = False
        retry_seconds = 0

        for attempt in range(2):
            try:
                result = self.executor.invoke({"input": input_text})
                output = (result.get("output") or "").strip()
                if output:
                    break
                print(f"⚠️  Tentativo {attempt + 1}: risposta vuota, ritento tra 3s...")
                time.sleep(3)
            except Exception as e:
                err_str = str(e)
                # Riconoscimento specifico errore di rate limit / quota
                if "429" in err_str or "ResourceExhausted" in err_str or "quota" in err_str.lower():
                    rate_limited = True
                    match = re.search(r"retry in (\d+)", err_str)
                    if match:
                        retry_seconds = int(match.group(1))
                    print(f"⚠️  Rate limit Gemini raggiunto. Riprova tra {retry_seconds}s.")
                    break
                print(f"⚠️  Tentativo {attempt + 1}: errore {err_str[:100]}")
                output = ""
                time.sleep(3)


        if rate_limited:
            wait_msg = f"tra circa {retry_seconds} secondi" if retry_seconds else "tra qualche secondo"
            output = (
                f"⏱️ In questo momento stiamo gestendo molte richieste contemporaneamente. "
                f"Riprova {wait_msg}."
            )
            source = "RATE_LIMIT"
            return {
                "risposta": output,
                "escalation": False,
                "source": source,
                "tools_used": [],
                "rate_limited": True,
            }

        # Se vuoto dopo retry → escalation
        if not output:
            output = (
                "Ci scusiamo, in questo momento non riusciamo a generare una risposta. "
                "Stiamo trasferendo la richiesta a un Customer Success Consultant."
            )
            _last_tools_used.append("ESCALATION")

        tools_used = get_tools_used()
        is_escalation = "ESCALATION" in tools_used or "ESCALATION_RICHIESTA" in output

        if is_escalation:
            source = "ESCALATION"
        elif "CUSTOMER" in tools_used:
            source = "CUSTOMER"
        elif "RAG" in tools_used:
            source = "RAG"
        else:
            source = "GENERAL"

        return {
            "risposta": output.replace("ESCALATION_RICHIESTA:", "").strip(),
            "escalation": is_escalation,
            "source": source,
            "tools_used": tools_used,
            "rate_limited": False,
        }


# ─────────────────────────────────────────────
# CLI per test
# ─────────────────────────────────────────────
if __name__ == "__main__":
    agent = CustomerSupportAgent()
    print("AI Support Agent — digita 'esci' per uscire\n")
    cliente = input("Nome cliente (es. 'Mario Rossi', invio per saltare): ").strip() or None
    if cliente:
        agent.set_active_client(cliente)
    while True:
        q = input("\nTu: ").strip()
        if q.lower() in ("esci", "exit", "quit"):
            break
        resp = agent.ask(q)
        print(f"\nAI Support Agent [{resp['source']}]: {resp['risposta']}")
        if resp["escalation"]:
            print("⚠️  [Escalation attivata — verrà contattato un consulente]")
